from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt


ITEM_RE = re.compile(
    r"\\item\s*\[\s*\\textbf\s*\{\s*(\d+)\s*\.?\s*\}\s*\](.*?)(?=\\item\s*\[\s*\\textbf\s*\{|\s*\\end\s*\{enumerate\}|\Z)",
    re.IGNORECASE | re.DOTALL,
)
TAG_RE = re.compile(r"\[\[\s*([A-Za-z_]+)\s*=\s*(.*?)\s*\]\]", re.IGNORECASE)
ENV_RE = re.compile(r"\\(?:begin|end)\s*\{(?:enumerate|document|itemize)\}", re.IGNORECASE)


def read_text_any_encoding(path: Path) -> str:
    raw = path.read_bytes()
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("latin-1", errors="replace")


def repair_mojibake(text: str) -> str:
    value = str(text or "")
    markers = ("Ã", "Â", "â")
    if any(marker in value for marker in markers):
        try:
            repaired = value.encode("latin-1", errors="strict").decode("utf-8", errors="strict")
            if sum(repaired.count(marker) for marker in markers) < sum(value.count(marker) for marker in markers):
                value = repaired
        except Exception:
            pass
    return value


def normalize_latex_text(text: str) -> str:
    value = repair_mojibake(text)
    value = ENV_RE.sub(" ", value)
    value = value.replace("\u00a3", "\n")
    value = value.replace("\u00e6", "\n")
    value = value.replace("\u00c2\u00a3", "\n")
    value = value.replace("\u00c3\u00a6", "\n")
    value = value.replace(r"\\ ", "\n")
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\n\s+", "\n", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def extract_tags(body: str) -> tuple[dict[str, str], str]:
    tags: dict[str, str] = {}

    def repl(match: re.Match[str]) -> str:
        key = match.group(1).strip().lower()
        val = repair_mojibake(match.group(2).strip())
        tags[key] = val
        return " "

    clean = TAG_RE.sub(repl, body)
    return tags, clean


def split_problem_text(body: str) -> tuple[str, list[str]]:
    text = normalize_latex_text(body)
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    statement: list[str] = []
    options: list[str] = []
    for line in lines:
        if re.match(r"^[A-E]\)", line, flags=re.IGNORECASE):
            options.append(line)
        else:
            statement.append(line)
    return " ".join(statement).strip(), options


def parse_tex_items(tex_text: str) -> list[dict[str, object]]:
    source = repair_mojibake(tex_text)
    items: list[dict[str, object]] = []
    for match in ITEM_RE.finditer(source):
        number = int(match.group(1))
        tags, clean_body = extract_tags(match.group(2) or "")
        statement, options = split_problem_text(clean_body)
        items.append(
            {
                "number": number,
                "statement": statement,
                "options": options,
                "tags": tags,
            }
        )
    if items:
        return items
    fallback = normalize_latex_text(source)
    return [{"number": 1, "statement": fallback, "options": [], "tags": {}}] if fallback else []


def apply_base_styles(doc: Document, style: str) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11 if style != "compact" else 10)
    for section in doc.sections:
        section.top_margin = Pt(42)
        section.bottom_margin = Pt(42)
        section.left_margin = Pt(48)
        section.right_margin = Pt(48)


def add_problem(doc: Document, item: dict[str, object]) -> None:
    number = int(item.get("number") or 0)
    tags = item.get("tags") if isinstance(item.get("tags"), dict) else {}
    statement = str(item.get("statement") or "").strip()
    options = [str(option).strip() for option in item.get("options") or [] if str(option).strip()]

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{number}. ")
    run.bold = True
    p.add_run(statement)

    if options:
        for option in options:
            op = doc.add_paragraph(option)
            op.paragraph_format.left_indent = Pt(18)
            op.paragraph_format.space_after = Pt(1)

    meta_parts = []
    for key in ("curso", "tema", "subtema", "clave", "estado"):
        value = str(tags.get(key, "")).strip()
        if value:
            meta_parts.append(f"{key}: {value}")
    if meta_parts:
        meta = doc.add_paragraph(" | ".join(meta_parts))
        meta.style = doc.styles["Intense Quote"] if "Intense Quote" in doc.styles else doc.styles["Normal"]
        meta.paragraph_format.space_after = Pt(8)


def convert_tex_to_docx(input_tex: Path, output_docx: Path, *, style: str = "standard") -> Path:
    tex_text = read_text_any_encoding(input_tex)
    items = parse_tex_items(tex_text)

    doc = Document()
    apply_base_styles(doc, style)
    title = input_tex.stem.replace("_", " ").strip() or "Practica"
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Fuente: {input_tex.name}")
    doc.add_paragraph(f"Total de problemas: {len(items)}")
    doc.add_paragraph("")

    for item in items:
        add_problem(doc, item)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))
    return output_docx


def main() -> int:
    parser = argparse.ArgumentParser(description="Convierte fuente LaTeX de Auditor-IA a Word.")
    parser.add_argument("input_tex", type=Path)
    parser.add_argument("output_docx", type=Path)
    parser.add_argument("--style", default="standard", choices=("standard", "compact"))
    parser.add_argument("--template", type=Path, default=None)
    parser.add_argument("--images-dir", type=Path, default=None)
    args = parser.parse_args()

    if not args.input_tex.exists():
        raise SystemExit(f"No existe el archivo .tex: {args.input_tex}")
    output = convert_tex_to_docx(args.input_tex, args.output_docx, style=args.style)
    print(f"Word generado en: {output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
