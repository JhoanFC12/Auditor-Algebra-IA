from __future__ import annotations

import re
import unicodedata
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional

from database.connection import DatabaseManager

try:
    from docx import Document
    from docx.shared import RGBColor
except Exception:  # pragma: no cover - optional dependency handled at runtime
    Document = None  # type: ignore[assignment]
    RGBColor = None  # type: ignore[assignment]


TAG_META_RE = re.compile(r"\[\[\s*(?:curso|tema|subtema|clave|estado|examen|solucion(?:ario)?)\s*=\s*.*?\s*\]\]", re.IGNORECASE)
ESTADO_TAG_RE = re.compile(r"\[\[\s*estado\s*=\s*([^\]]+?)\s*\]\]", re.IGNORECASE)
IMG_MARKER_RE = re.compile(r"\[\[\s*imagen\s*=\s*.*?\s*\]\]", re.IGNORECASE)
ITEM_HEADER_RE = re.compile(
    r"^\s*(\\item\s*\[\s*\\textbf\{\s*\d+\.?\s*\}\s*\])\s*(.*)$",
    re.IGNORECASE | re.DOTALL,
)
OPTION_LABEL_RE = re.compile(r"(?<![A-Za-z0-9])([A-E])\)", re.IGNORECASE)


class PracticeBuilderController:
    def __init__(self, db_manager: DatabaseManager | None = None):
        self.db = db_manager or DatabaseManager()
        self._schema_cache: dict[str, dict[str, object]] = {}

    def listar_bases_datos(self) -> List[str]:
        return self.db.listar_bases_datos()

    def describir_fuente(self, db_name: str) -> Dict[str, object]:
        return dict(self._schema_info(db_name))

    def normalizar_curso(self, value: object) -> str:
        return self._canonical_course_value(value)

    def clave_texto_normalizada(self, value: object) -> str:
        return self._normalized_text_key(value)

    def listar_cursos(self, db_name: str) -> List[str]:
        schema = self._schema_info(db_name)
        values: list[str] = []
        if schema["has_temas_table"]:
            values.extend(
                self._query_distinct(
                db_name,
                """
                SELECT DISTINCT COALESCE(t.area,'')
                FROM temas t
                WHERE COALESCE(t.area,'') <> ''
                ORDER BY COALESCE(t.area,'') ASC;
                """,
                )
            )
        if schema["course_column"]:
            values.extend(
                self._query_distinct(
                db_name,
                f"""
                SELECT DISTINCT TRIM(CAST(p.{schema["course_column"]} AS text))
                FROM problemas p
                WHERE COALESCE(TRIM(CAST(p.{schema["course_column"]} AS text)), '') <> ''
                ORDER BY TRIM(CAST(p.{schema["course_column"]} AS text)) ASC;
                """,
                )
            )
        return self._normalize_course_values(values)

    def listar_temas(self, db_name: str, *, curso: str = "") -> List[Dict[str, object]]:
        schema = self._schema_info(db_name)
        curso = self._canonical_course_value(curso)
        result: list[dict[str, object]] = []
        seen: set[tuple[str, str]] = set()

        def append_topic(topic_id: object, nombre: object, area: object = "") -> None:
            clean_name = str(nombre or "").strip()
            clean_area = self._canonical_course_value(area)
            if not clean_name:
                return
            key = (self._normalized_text_key(clean_area), self._normalized_text_key(clean_name))
            if key in seen:
                return
            seen.add(key)
            result.append({"id": topic_id, "nombre": clean_name, "curso": clean_area})

        if schema["has_temas_table"]:
            if curso:
                normalized_area_expr = self._sql_normalized_text("COALESCE(t.area,'')")
                rows = self._query_rows(
                    db_name,
                    f"""
                    SELECT t.id, COALESCE(t.nombre,''), COALESCE(t.area,'')
                    FROM temas t
                    WHERE {normalized_area_expr} = %s
                    ORDER BY COALESCE(t.nombre,'') ASC;
                    """,
                    (self._normalized_text_key(curso),),
                )
            else:
                rows = self._query_rows(
                    db_name,
                    """
                    SELECT t.id, COALESCE(t.nombre,''), COALESCE(t.area,'')
                    FROM temas t
                    ORDER BY COALESCE(t.area,''), COALESCE(t.nombre,'') ASC;
                    """,
                )
            for r in rows:
                append_topic(self._catalog_ref("topic", r[0], r[1]), r[1], r[2])

        topic_column = str(schema["topic_column"] or "")
        if not topic_column:
            return result

        clauses = [f"COALESCE(TRIM(CAST(p.{topic_column} AS text)), '') <> ''"]
        params: list[object] = []
        course_column = str(schema["course_column"] or "")
        if curso and course_column:
            normalized_course_expr = self._sql_normalized_text(f"p.{course_column}")
            clauses.append(f"{normalized_course_expr} = %s")
            params.append(self._normalized_text_key(curso))
        rows = self._query_rows(
            db_name,
            f"""
            SELECT DISTINCT TRIM(CAST(p.{topic_column} AS text))
            FROM problemas p
            WHERE {' AND '.join(clauses)}
            ORDER BY TRIM(CAST(p.{topic_column} AS text)) ASC;
            """,
            tuple(params),
        )
        names = [str(r[0] or "").strip() for r in rows if str(r[0] or "").strip()]
        for name in self._normalize_text_values(names):
            append_topic(self._direct_ref("topic", name), name, curso)
        return result

    def listar_subtemas(self, db_name: str, *, tema_id: Optional[object]) -> List[Dict[str, object]]:
        schema = self._schema_info(db_name)
        result: list[dict[str, object]] = []
        seen: set[str] = set()

        def append_subtopic(subtopic_id: object, nombre: object) -> None:
            clean_name = str(nombre or "").strip()
            if not clean_name:
                return
            key = self._normalized_text_key(clean_name)
            if key in seen:
                return
            seen.add(key)
            result.append({"id": subtopic_id, "nombre": clean_name})

        topic_ref = self._parse_meta_ref(tema_id)
        if schema["has_subtemas_table"]:
            if topic_ref["id"] is not None:
                rows = self._query_rows(
                    db_name,
                    """
                    SELECT s.id, COALESCE(s.nombre,'')
                    FROM subtemas s
                    WHERE s.tema_id = %s
                    ORDER BY COALESCE(s.nombre,'') ASC;
                    """,
                    (int(topic_ref["id"]),),
                )
                for r in rows:
                    append_subtopic(self._catalog_ref("subtopic", r[0], r[1]), r[1])

        subtopic_column = str(schema["subtopic_column"] or "")
        if not subtopic_column:
            return result
        clauses = [f"COALESCE(TRIM(CAST(p.{subtopic_column} AS text)), '') <> ''"]
        params: list[object] = []
        topic_column = str(schema["topic_column"] or "")
        if topic_ref["text"] and topic_column:
            normalized_topic_expr = self._sql_normalized_text(f"p.{topic_column}")
            clauses.append(f"{normalized_topic_expr} = %s")
            params.append(self._normalized_text_key(str(topic_ref["text"])))
        elif topic_ref["id"] is not None and schema.get("tema_id_column"):
            clauses.append("p.tema_id = %s")
            params.append(int(topic_ref["id"]))
        rows = self._query_rows(
            db_name,
            f"""
            SELECT DISTINCT TRIM(CAST(p.{subtopic_column} AS text))
            FROM problemas p
            WHERE {' AND '.join(clauses)}
            ORDER BY TRIM(CAST(p.{subtopic_column} AS text)) ASC;
            """,
            tuple(params),
        )
        names = [str(r[0] or "").strip() for r in rows if str(r[0] or "").strip()]
        for name in self._normalize_text_values(names):
            append_subtopic(self._direct_ref("subtopic", name), name)
        return result

    def listar_autores(
        self,
        db_name: str,
        *,
        curso: str = "",
        tema_id: Optional[object] = None,
        subtema_id: Optional[object] = None,
        editorial: str = "",
    ) -> List[str]:
        return self._list_problem_values(
            db_name,
            "author_column",
            curso=curso,
            tema_id=tema_id,
            subtema_id=subtema_id,
            editorial=editorial,
        )

    def listar_editoriales(
        self,
        db_name: str,
        *,
        curso: str = "",
        tema_id: Optional[object] = None,
        subtema_id: Optional[object] = None,
        autor: str = "",
    ) -> List[str]:
        return self._list_problem_values(
            db_name,
            "editorial_column",
            curso=curso,
            tema_id=tema_id,
            subtema_id=subtema_id,
            autor=autor,
        )

    def contar_problemas(
        self,
        db_name: str,
        *,
        curso: str = "",
        tema_id: Optional[object] = None,
        subtema_id: Optional[object] = None,
        autor: str = "",
        editorial: str = "",
        estado: str = "Todos",
        clave: str = "Todos",
    ) -> int:
        schema = self._schema_info(db_name)
        join_sql, where_sql, params = self._build_filters(
            schema,
            curso=curso,
            tema_id=tema_id,
            subtema_id=subtema_id,
            autor=autor,
            editorial=editorial,
            estado=estado,
            clave=clave,
        )
        rows = self._query_rows(
            db_name,
            f"SELECT COUNT(*)::int FROM problemas p {join_sql} {where_sql};",
            tuple(params),
        )
        return int(rows[0][0] or 0) if rows else 0

    def obtener_problemas(
        self,
        db_name: str,
        *,
        cantidad: int,
        curso: str = "",
        tema_id: Optional[object] = None,
        subtema_id: Optional[object] = None,
        autor: str = "",
        editorial: str = "",
        estado: str = "Todos",
        clave: str = "Todos",
        aleatorio: bool = True,
    ) -> List[Dict[str, object]]:
        cantidad = max(int(cantidad or 0), 1)
        schema = self._schema_info(db_name)
        join_sql, where_sql, params = self._build_filters(
            schema,
            curso=curso,
            tema_id=tema_id,
            subtema_id=subtema_id,
            autor=autor,
            editorial=editorial,
            estado=estado,
            clave=clave,
        )
        order_sql = "ORDER BY random()" if aleatorio else self._natural_problem_order_sql(schema)
        rows = self._query_rows(
            db_name,
            f"""
            SELECT
                {self._problem_select_columns_sql(schema)}
            FROM problemas p
            {join_sql}
            {where_sql}
            {order_sql}
            LIMIT %s;
            """,
            (*params, cantidad),
        )
        return [self._problem_row_to_dict(r) for r in rows]

    def obtener_problema_por_id(self, db_name: str, *, problem_id: int) -> Dict[str, object] | None:
        schema = self._schema_info(db_name)
        join_sql = self._compose_join_sql(schema)
        rows = self._query_rows(
            db_name,
            f"""
            SELECT
                {self._problem_select_columns_sql(schema)}
            FROM problemas p
            {join_sql}
            WHERE p.id = %s
            LIMIT 1;
            """,
            (int(problem_id),),
        )
        if not rows:
            return None
        return self._problem_row_to_dict(rows[0])

    def actualizar_enunciado_problema(
        self,
        db_name: str,
        *,
        problem_id: int,
        enunciado_latex: str,
        updated_by: str = "modulo7-mathjax",
    ) -> Dict[str, object]:
        schema = self._schema_info(db_name)
        problem_cols = set(str(x) for x in schema.get("problem_columns", []))
        updates = ["enunciado_latex = %s"]
        params: list[object] = [self._strip_metadata_tags(str(enunciado_latex or ""))]
        if "updated_at" in problem_cols:
            updates.append("updated_at = NOW()")
        if "updated_by" in problem_cols:
            updates.append("updated_by = %s")
            params.append(str(updated_by or "").strip() or "modulo7-mathjax")
        if "revision_version" in problem_cols:
            updates.append("revision_version = COALESCE(revision_version, 0) + 1")

        conn = self.db.get_connection(db_name)
        try:
            cur = conn.cursor()
            cur.execute(
                f"""
                UPDATE problemas
                SET {", ".join(updates)}
                WHERE id = %s
                RETURNING id;
                """,
                (*params, int(problem_id)),
            )
            row = cur.fetchone()
            if not row:
                raise RuntimeError(f"No existe el problema ID {int(problem_id)}.")
            conn.commit()
        except Exception:
            conn.rollback()
            raise
        finally:
            conn.close()

        refreshed = self.obtener_problema_por_id(db_name, problem_id=int(problem_id))
        if refreshed is None:
            raise RuntimeError(f"No se pudo recargar el problema ID {int(problem_id)} despues de guardar.")
        return refreshed

    def obtener_problemas_por_ids(self, db_name: str, *, problem_ids: List[int]) -> List[Dict[str, object]]:
        ordered_ids: list[int] = []
        seen: set[int] = set()
        for raw_id in problem_ids:
            try:
                pid = int(raw_id)
            except Exception:
                continue
            if pid <= 0 or pid in seen:
                continue
            seen.add(pid)
            ordered_ids.append(pid)
        if not ordered_ids:
            return []

        schema = self._schema_info(db_name)
        problem_cols = set(str(x) for x in schema.get("problem_columns", []))
        join_sql = self._compose_join_sql(schema)
        response_expr = self._response_select_expr(schema)
        course_expr = self._meta_select_expr(schema, "course_column", "COALESCE(t.area,'')", use_catalog=True)
        topic_expr = self._meta_select_expr(schema, "topic_column", "COALESCE(t.nombre,'')", use_catalog=True)
        subtopic_expr = self._meta_select_expr(schema, "subtopic_column", "COALESCE(s.nombre,'')", use_catalog=True)
        author_expr = self._meta_select_expr(schema, "author_column", "COALESCE(l.autor,'')", allow_fallback=True)
        editorial_expr = self._meta_select_expr(
            schema,
            "editorial_column",
            "COALESCE(l.editorial,'')",
            allow_fallback=True,
        )
        math_consistency_column = str(schema.get("math_consistency_column") or "")
        source_column = str(schema.get("source_file_column") or "")
        folder_column = str(schema.get("folder_column") or "")
        images_column = str(schema.get("images_column") or "")
        key_flag_column = str(schema.get("key_flag_column") or "")
        problem_type_column = str(schema.get("problem_type_column") or "")
        rows = self._query_rows(
            db_name,
            f"""
            SELECT
                p.id,
                p.numero_original,
                COALESCE(p.enunciado_latex,''),
                {response_expr},
                {course_expr},
                {topic_expr},
                {subtopic_expr},
                {author_expr},
                {editorial_expr},
                {"COALESCE(l.pdf_path,'')" if self._needs_book_join(schema) else "''"},
                {f"COALESCE(p.{source_column},'')" if source_column else "''"},
                {f"COALESCE(p.{folder_column},'')" if folder_column else "''"},
                {f"COALESCE(p.{images_column}, ARRAY[]::text[])" if images_column else "ARRAY[]::text[]"},
                {f"COALESCE(p.{math_consistency_column},'')" if math_consistency_column else "'Sin revisar'"},
                {"COALESCE(p.codigo_instancia,'')" if "codigo_instancia" in problem_cols else "''"},
                {"COALESCE(p.libro_codigo,'')" if "libro_codigo" in problem_cols else "''"},
                {f"COALESCE(p.{key_flag_column}, FALSE)" if key_flag_column else "FALSE"},
                {"COALESCE(p.tiene_solucion, FALSE)" if "tiene_solucion" in problem_cols else "FALSE"},
                {f"COALESCE(p.{problem_type_column},'')" if problem_type_column else "''"}
            FROM problemas p
            {join_sql}
            WHERE p.id = ANY(%s)
            ORDER BY p.id ASC;
            """,
            (ordered_ids,),
        )
        items: list[dict[str, object]] = []
        for row in rows:
            raw_images = row[12] if len(row) > 12 else []
            images: list[str] = []
            if isinstance(raw_images, (list, tuple)):
                for value in raw_images:
                    clean = str(value or "").strip()
                    if clean:
                        images.append(clean)
            item_text = str(row[2] or "")
            respuesta = str(row[3] or "")
            explicit_type = str(row[18] if len(row) > 18 else "")
            items.append(
                {
                    "id": int(row[0]),
                    "numero_original": int(row[1] or 0),
                    "enunciado_latex": item_text,
                    "respuesta_correcta": respuesta,
                    "curso": self._canonical_course_value(row[4]),
                    "tema": str(row[5] or ""),
                    "subtema": str(row[6] or ""),
                    "autor": str(row[7] or ""),
                    "editorial": str(row[8] or ""),
                    "pdf_path": str(row[9] or ""),
                    "archivo_origen": str(row[10] or ""),
                    "ruta_carpeta": str(row[11] or ""),
                    "imagenes": images,
                    "consistencia_matematica": str(row[13] or "") or "Sin revisar",
                    "codigo_instancia": str(row[14] or ""),
                    "libro_codigo": str(row[15] or ""),
                    "tiene_clave": bool(row[16]),
                    "tiene_solucion": bool(row[17]),
                    "tipo_problema": self._infer_problem_type(item_text, respuesta, explicit_type),
                }
            )
        by_id = {int(item["id"]): item for item in items}
        return [by_id[pid] for pid in ordered_ids if pid in by_id]

    def actualizar_problema_desde_editor(
        self,
        db_name: str,
        *,
        problem_id: int,
        numero_original: int,
        enunciado_latex: str,
        respuesta_correcta: str,
        curso: str,
        tema: str,
        subtema: str,
        consistencia_matematica: str,
        updated_by: str = "modulo6-json-editor",
    ) -> Dict[str, object]:
        schema = self._schema_info(db_name)
        problem_cols = set(str(x) for x in schema.get("problem_columns", []))
        key_flag_column = str(schema.get("key_flag_column") or "")
        math_consistency_column = str(schema.get("math_consistency_column") or "")

        clean_key = str(respuesta_correcta or "").strip().upper()
        clean_math_consistency = str(consistencia_matematica or "").strip() or "Sin revisar"
        updates = [
            "numero_original = %s",
            "enunciado_latex = %s",
            "respuesta_correcta = %s",
            "curso = %s",
            "tema = %s",
            "subtema = %s",
        ]
        params: list[object] = [
            max(int(numero_original or 0), 1),
            self._strip_metadata_tags(str(enunciado_latex or "")),
            clean_key,
            str(curso or "").strip(),
            str(tema or "").strip(),
            str(subtema or "").strip(),
        ]
        if math_consistency_column:
            updates.append(f"{math_consistency_column} = %s")
            params.append(clean_math_consistency)
        if key_flag_column:
            updates.append(f"{key_flag_column} = %s")
            params.append(bool(clean_key))
        if "updated_at" in problem_cols:
            updates.append("updated_at = NOW()")
        if "updated_by" in problem_cols:
            updates.append("updated_by = %s")
            params.append(str(updated_by or "").strip() or "modulo6-json-editor")
        if "revision_version" in problem_cols:
            updates.append("revision_version = COALESCE(revision_version, 0) + 1")

        conn = self.db.get_connection(db_name)
        try:
            cur = conn.cursor()
            cur.execute(
                f"""
                UPDATE problemas
                SET {", ".join(updates)}
                WHERE id = %s
                RETURNING id;
                """,
                (*params, int(problem_id)),
            )
            row = cur.fetchone()
            if not row:
                raise RuntimeError(f"No existe el problema ID {int(problem_id)}.")
            conn.commit()
        except Exception:
            conn.rollback()
            raise
        finally:
            conn.close()

        refreshed_rows = self.obtener_problemas_por_ids(db_name, problem_ids=[int(problem_id)])
        if not refreshed_rows:
            raise RuntimeError(f"No se pudo recargar el problema ID {int(problem_id)} despues de guardar.")
        return refreshed_rows[0]

    def exportar_practica_word(
        self,
        *,
        output_path: Path,
        titulo: str,
        problemas: List[Dict[str, object]],
        incluir_clave_inline: bool = False,
        incluir_clave_final: bool = True,
    ) -> Path:
        if Document is None or RGBColor is None:
            raise RuntimeError("Falta dependencia python-docx. Instala requirements.txt en tu entorno activo.")

        doc = Document()
        doc.add_heading(titulo.strip() or "Practica", level=1)
        doc.add_paragraph(f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"Total de problemas: {len(problemas)}")
        doc.add_paragraph("")

        for idx, item in enumerate(problemas, start=1):
            enunciado = self._limpiar_enunciado_para_word(str(item.get("enunciado_latex") or ""))
            p = doc.add_paragraph()
            p.add_run(f"{idx}. ")
            p.add_run(enunciado)
            clave = (item.get("respuesta_correcta") or "").strip().upper()
            if incluir_clave_inline and clave:
                run = p.add_run(f"   [Clave: {clave}]")
                run.font.color.rgb = RGBColor(220, 38, 38)

        if incluir_clave_final:
            doc.add_page_break()
            doc.add_heading("Clave de respuestas", level=2)
            for idx, item in enumerate(problemas, start=1):
                clave = (item.get("respuesta_correcta") or "").strip().upper()
                if not clave:
                    continue
                line = doc.add_paragraph()
                line.add_run(f"{idx}) ")
                run = line.add_run(clave)
                run.font.color.rgb = RGBColor(220, 38, 38)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return output_path

    def renderizar_practica_latex(
        self,
        *,
        titulo: str,
        problemas: List[Dict[str, object]],
        curso: str = "",
        tema: str = "",
        subtema: str = "",
        incluir_clave_final: bool = True,
    ) -> str:
        template_path = Path(__file__).resolve().parent / "templates" / "practica_base.tex"
        if not template_path.exists():
            raise RuntimeError(f"No existe plantilla base: {template_path}")

        template = template_path.read_text(encoding="utf-8")
        items = [
            self._normalizar_item_para_latex(str(p.get("enunciado_latex") or ""), idx)
            for idx, p in enumerate(problemas, start=1)
        ]
        body = "\n\n".join([x for x in items if x.strip()]).strip() or "% (sin problemas)"

        key_lines: List[str] = []
        if incluir_clave_final:
            for idx, p in enumerate(problemas, start=1):
                clave = str(p.get("respuesta_correcta") or "").strip().upper()
                if clave:
                    key_lines.append(f"{idx}) {self._escape_latex_text(clave)}")
        key_block = "\n".join(key_lines).strip() or "Sin clave registrada."

        rendered = template
        rendered = rendered.replace("{{TITLE}}", self._escape_latex_text((titulo or "").strip() or "Practica"))
        rendered = rendered.replace("{{DATE}}", datetime.now().strftime("%Y-%m-%d %H:%M"))
        rendered = rendered.replace("{{COURSE}}", self._escape_latex_text((curso or "").strip() or "Todos"))
        rendered = rendered.replace("{{TOPIC}}", self._escape_latex_text((tema or "").strip() or "Todos"))
        rendered = rendered.replace("{{SUBTOPIC}}", self._escape_latex_text((subtema or "").strip() or "Todos"))
        rendered = rendered.replace("{{TOTAL}}", str(len(problemas)))
        rendered = rendered.replace("{{PROBLEMS}}", body)
        rendered = rendered.replace("{{ANSWER_KEY}}", key_block)
        rendered = rendered.replace("{{SHOW_ANSWER_KEY}}", "true" if incluir_clave_final else "false")
        return rendered

    def exportar_practica_latex(
        self,
        *,
        output_path: Path,
        titulo: str,
        problemas: List[Dict[str, object]],
        curso: str = "",
        tema: str = "",
        subtema: str = "",
        incluir_clave_final: bool = True,
    ) -> Path:
        content = self.renderizar_practica_latex(
            titulo=titulo,
            problemas=problemas,
            curso=curso,
            tema=tema,
            subtema=subtema,
            incluir_clave_final=incluir_clave_final,
        )
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(content, encoding="utf-8")
        return output_path

    def renderizar_fuente_latex_para_word(self, *, problemas: List[Dict[str, object]]) -> str:
        blocks: list[str] = []
        for problema in problemas:
            raw = str(problema.get("enunciado_latex") or "").strip()
            if not raw:
                continue
            clave = str(problema.get("respuesta_correcta") or "").strip().upper()
            if clave and not re.search(r"\[\[\s*clave\s*=", raw, flags=re.IGNORECASE):
                raw = f"{raw} [[clave={clave}]]"
            estado_bruto = str(problema.get("consistencia_matematica") or "").strip()
            estado_tag = self._normalizar_estado_tag(estado_bruto)
            if estado_tag and not ESTADO_TAG_RE.search(raw):
                raw = f"{raw} [[Estado={estado_tag}]]"
            blocks.append(raw)
        source_text = "\n".join(blocks).strip()
        return self._ensure_enumerate_wrapper(source_text)

    def _normalizar_estado_tag(self, estado: str) -> str:
        valor = str(estado or "").strip().lower()
        if not valor:
            return "sin_revisar"
        if valor in {"consistente", "bien planteado", "bien_planteado"}:
            return "consistente"
        if valor in {"inconsistente", "mal planteado", "mal_planteado", "ambiguo", "ambigua"}:
            return "inconsistente"
        if valor in {"sin revisar", "sin_revisar", "pendiente revision", "pendiente revisión"}:
            return "sin_revisar"
        return valor.replace(" ", "_")

    def _strip_metadata_tags(self, text: str) -> str:
        clean = TAG_META_RE.sub(" ", str(text or ""))
        clean = re.sub(r"[ \t]{2,}", " ", clean)
        clean = re.sub(r"\n{3,}", "\n\n", clean)
        return clean.strip()

    def exportar_fuente_latex_para_word(
        self,
        *,
        output_path: Path,
        problemas: List[Dict[str, object]],
    ) -> Path:
        content = self.renderizar_fuente_latex_para_word(problemas=problemas)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(content + "\n", encoding="utf-8")
        return output_path

    def _schema_info(self, db_name: str) -> Dict[str, object]:
        cached = self._schema_cache.get(db_name)
        if cached is not None:
            return cached

        conn = self.db.get_connection(db_name)
        try:
            cur = conn.cursor()
            cur.execute(
                """
                SELECT table_name
                FROM information_schema.tables
                WHERE table_schema='public' AND table_name IN ('problemas','temas','subtemas','libros_escaneo','origenes','problema_origen');
                """
            )
            tables = {str(r[0]) for r in cur.fetchall()}
            cur.execute(
                """
                SELECT table_name, column_name
                FROM information_schema.columns
                WHERE table_schema='public' AND table_name IN ('problemas','temas','subtemas','libros_escaneo','origenes','problema_origen');
                """
            )
            columns: dict[str, set[str]] = {
                "problemas": set(),
                "temas": set(),
                "subtemas": set(),
                "libros_escaneo": set(),
                "origenes": set(),
                "problema_origen": set(),
            }
            for table_name, column_name in cur.fetchall():
                columns.setdefault(str(table_name), set()).add(str(column_name))
        finally:
            conn.close()

        problem_cols = columns.get("problemas", set())
        book_cols = columns.get("libros_escaneo", set())
        info: Dict[str, object] = {
            "has_problemas_table": "problemas" in tables,
            "has_temas_table": "temas" in tables and {"id", "nombre"}.issubset(columns.get("temas", set())),
            "has_subtemas_table": "subtemas" in tables and {"id", "nombre"}.issubset(columns.get("subtemas", set())),
            "has_books_table": "libros_escaneo" in tables and {"id", "codigo"}.issubset(book_cols),
            "has_origins_table": "origenes" in tables
            and "problema_origen" in tables
            and {"id", "nombre", "tipo_origen"}.issubset(columns.get("origenes", set()))
            and {"problema_id", "origen_id"}.issubset(columns.get("problema_origen", set())),
            "problem_columns": sorted(problem_cols),
            "course_column": "curso" if "curso" in problem_cols else None,
            "topic_column": "tema" if "tema" in problem_cols else None,
            "subtopic_column": "subtema" if "subtema" in problem_cols else None,
            "author_column": "autor" if "autor" in problem_cols else None,
            "editorial_column": "editorial" if "editorial" in problem_cols else None,
            "response_column": "respuesta_correcta" if "respuesta_correcta" in problem_cols else ("respuesta" if "respuesta" in problem_cols else None),
            "key_flag_column": "tiene_clave" if "tiene_clave" in problem_cols else None,
            "problem_type_column": "tipo_problema" if "tipo_problema" in problem_cols else None,
            "math_consistency_column": "consistencia_matematica" if "consistencia_matematica" in problem_cols else None,
            "tema_id_column": "tema_id" if "tema_id" in problem_cols else None,
            "subtema_id_column": "subtema_id" if "subtema_id" in problem_cols else None,
            "book_id_column": "libro_id" if "libro_id" in problem_cols else None,
            "book_code_column": "libro_codigo" if "libro_codigo" in problem_cols else None,
            "source_file_column": "archivo_origen" if "archivo_origen" in problem_cols else None,
            "images_column": "imagenes" if "imagenes" in problem_cols else None,
            "folder_column": "ruta_carpeta" if "ruta_carpeta" in problem_cols else None,
        }
        info["uses_catalog"] = bool(info["has_temas_table"]) and bool(info["tema_id_column"])
        self._schema_cache[db_name] = info
        return info

    def _query_rows(self, db_name: str, sql: str, params: tuple[object, ...] = ()) -> list[tuple]:
        conn = self.db.get_connection(db_name)
        try:
            cur = conn.cursor()
            cur.execute(sql, params)
            return list(cur.fetchall())
        finally:
            conn.close()

    def _query_distinct(self, db_name: str, sql: str, params: tuple[object, ...] = ()) -> list[str]:
        rows = self._query_rows(db_name, sql, params)
        return [str(r[0]) for r in rows if str(r[0] or "").strip()]

    def _normalize_course_values(self, values: List[str]) -> List[str]:
        deduped: dict[str, str] = {}
        for raw_value in values:
            canonical = self._canonical_course_value(raw_value)
            if not canonical:
                continue
            key = self._normalized_text_key(canonical)
            deduped[key] = self._prefer_display_variant(deduped.get(key, ""), canonical)
        return sorted(deduped.values(), key=self._normalized_text_key)

    def _normalize_text_values(self, values: List[str]) -> List[str]:
        deduped: dict[str, str] = {}
        for raw_value in values:
            cleaned = str(raw_value or "").strip()
            if not cleaned:
                continue
            key = self._normalized_text_key(cleaned)
            deduped[key] = self._prefer_display_variant(deduped.get(key, ""), cleaned)
        return sorted(deduped.values(), key=self._normalized_text_key)

    def _prefer_display_variant(self, current: str, candidate: str) -> str:
        if not current:
            return candidate
        current_score = self._display_variant_score(current)
        candidate_score = self._display_variant_score(candidate)
        if candidate_score > current_score:
            return candidate
        if candidate_score < current_score:
            return current
        return candidate if candidate < current else current

    def _display_variant_score(self, value: str) -> tuple[int, int, int]:
        text = str(value or "").strip()
        accent_score = sum(
            1 for ch in text if self._normalized_text_key(ch) != ch.casefold()
        )
        upper_score = sum(1 for ch in text if ch.isupper())
        return accent_score, len(text), -upper_score

    def _normalized_text_key(self, value: object) -> str:
        text = str(value or "").strip()
        if not text:
            return ""
        text = unicodedata.normalize("NFKD", text)
        text = "".join(ch for ch in text if not unicodedata.combining(ch))
        text = re.sub(r"\s+", " ", text)
        return text.casefold()

    def _direct_ref(self, kind: str, text: object) -> str:
        clean_kind = str(kind or "").strip()
        clean_text = str(text or "").strip()
        return f"direct:{clean_kind}:{clean_text}"

    def _catalog_ref(self, kind: str, raw_id: object, text: object = "") -> str:
        clean_kind = str(kind or "").strip()
        clean_id = str(raw_id or "").strip()
        clean_text = str(text or "").strip()
        return f"catalog:{clean_kind}:{clean_id}:{clean_text}"

    def _parse_meta_ref(self, value: object) -> dict[str, object]:
        raw = str(value or "").strip()
        parsed: dict[str, object] = {"source": "", "kind": "", "id": None, "text": ""}
        if not raw or raw == "Todos":
            return parsed

        if raw.startswith("catalog:"):
            parts = raw.split(":", 3)
            parsed["source"] = "catalog"
            parsed["kind"] = parts[1].strip() if len(parts) > 1 else ""
            raw_id = parts[2].strip() if len(parts) > 2 else ""
            parsed["text"] = parts[3].strip() if len(parts) > 3 else ""
            try:
                parsed["id"] = int(raw_id)
            except Exception:
                parsed["id"] = None
            return parsed

        if raw.startswith("direct:"):
            parts = raw.split(":", 2)
            parsed["source"] = "direct"
            parsed["kind"] = parts[1].strip() if len(parts) > 1 else ""
            parsed["text"] = parts[2].strip() if len(parts) > 2 else ""
            return parsed

        try:
            parsed["source"] = "catalog"
            parsed["id"] = int(raw)
        except Exception:
            parsed["source"] = "direct"
            parsed["text"] = raw
        return parsed

    def _canonical_course_value(self, value: object) -> str:
        text = str(value or "").strip()
        if not text:
            return ""
        normalized = self._normalized_text_key(text)
        if normalized == "trigonometria":
            return "Trigonometría"
        return text

    def _sql_normalized_text(self, expression: str) -> str:
        return (
            "LOWER(TRANSLATE(TRIM(CAST("
            + expression
            + " AS text)), 'ÁÉÍÓÚÜáéíóúü', 'AEIOUUaeiouu'))"
        )

    def _list_problem_values(
        self,
        db_name: str,
        field_key: str,
        *,
        curso: str = "",
        tema_id: Optional[object] = None,
        subtema_id: Optional[object] = None,
        autor: str = "",
        editorial: str = "",
    ) -> List[str]:
        schema = self._schema_info(db_name)
        field_expr = self._filter_field_expr(schema, field_key)
        if not field_expr:
            return []
        join_sql, where_sql, params = self._build_filters(
            schema,
            curso=curso,
            tema_id=tema_id,
            subtema_id=subtema_id,
            autor=autor if field_key != "author_column" else "",
            editorial=editorial if field_key != "editorial_column" else "",
            estado="Todos",
            clave="Todos",
        )
        rows = self._query_rows(
            db_name,
            f"""
            SELECT DISTINCT TRIM(CAST({field_expr} AS text))
            FROM problemas p
            {join_sql}
            {where_sql}
            {"AND" if where_sql else "WHERE"} COALESCE(TRIM(CAST({field_expr} AS text)), '') <> ''
            ORDER BY TRIM(CAST({field_expr} AS text)) ASC;
            """,
            tuple(params),
        )
        return self._normalize_text_values([str(r[0]) for r in rows if str(r[0] or "").strip()])

    def _response_select_expr(self, schema: Dict[str, object]) -> str:
        response_column = str(schema.get("response_column") or "")
        if response_column:
            return f"COALESCE(p.{response_column},'')"
        return "''"

    def _meta_select_expr(
        self,
        schema: Dict[str, object],
        key: str,
        fallback_expr: str,
        *,
        use_catalog: bool = False,
        allow_fallback: bool = False,
    ) -> str:
        column = str(schema.get(key) or "")
        if column:
            direct_expr = f"NULLIF(TRIM(CAST(p.{column} AS text)), '')"
            should_fallback = (use_catalog and schema.get("uses_catalog")) or allow_fallback
            if should_fallback and fallback_expr:
                return f"COALESCE({direct_expr}, {fallback_expr}, '')"
            return f"COALESCE(p.{column},'')"
        if use_catalog and schema.get("uses_catalog"):
            return fallback_expr
        if allow_fallback and fallback_expr:
            return fallback_expr
        return "''"

    def _filter_field_expr(self, schema: Dict[str, object], key: str) -> str:
        column = str(schema.get(key) or "")
        if column:
            return f"p.{column}"
        if key == "author_column" and self._needs_book_join(schema):
            return "l.autor"
        if key == "editorial_column" and self._needs_book_join(schema):
            return "l.editorial"
        return ""

    def _needs_book_join(self, schema: Dict[str, object]) -> bool:
        return bool(schema.get("has_books_table")) and bool(
            schema.get("book_id_column") or schema.get("book_code_column")
        )

    def _book_join_sql(self, schema: Dict[str, object]) -> str:
        if not self._needs_book_join(schema):
            return ""
        book_id_column = str(schema.get("book_id_column") or "")
        book_code_column = str(schema.get("book_code_column") or "")
        conditions: list[str] = []
        order_rules: list[str] = []
        if book_id_column:
            conditions.append(f"(p.{book_id_column} IS NOT NULL AND le.id = p.{book_id_column})")
            order_rules.append(f"CASE WHEN p.{book_id_column} IS NOT NULL AND le.id = p.{book_id_column} THEN 0 ELSE 1 END")
        if book_code_column:
            conditions.append(
                f"(COALESCE(TRIM(CAST(p.{book_code_column} AS text)), '') <> '' AND le.codigo = p.{book_code_column})"
            )
            order_rules.append(
                f"CASE WHEN COALESCE(TRIM(CAST(p.{book_code_column} AS text)), '') <> '' AND le.codigo = p.{book_code_column} THEN 0 ELSE 1 END"
            )
        if not conditions:
            return ""
        order_sql = ", ".join(order_rules + ["le.id"])
        where_sql = " OR ".join(f"({condition})" for condition in conditions)
        return f"""
            LEFT JOIN LATERAL (
                SELECT
                    le.id,
                    COALESCE(le.autor,'') AS autor,
                    COALESCE(le.editorial,'') AS editorial,
                    COALESCE(le.pdf_path,'') AS pdf_path
                FROM libros_escaneo le
                WHERE {where_sql}
                ORDER BY {order_sql}
                LIMIT 1
            ) l ON TRUE
        """

    def _compose_join_sql(self, schema: Dict[str, object]) -> str:
        joins: list[str] = []
        if schema.get("uses_catalog"):
            joins.append("LEFT JOIN temas t ON t.id = p.tema_id LEFT JOIN subtemas s ON s.id = p.subtema_id")
        book_join = self._book_join_sql(schema)
        if book_join.strip():
            joins.append(book_join.strip())
        if schema.get("has_origins_table"):
            joins.append(
                """
                LEFT JOIN LATERAL (
                    SELECT COALESCE(o.nombre, '') AS examen_nombre
                    FROM problema_origen po
                    JOIN origenes o ON o.id = po.origen_id
                    WHERE po.problema_id = p.id
                      AND o.tipo_origen = 'examen_admision'
                    ORDER BY po.orden NULLS LAST, po.id ASC
                    LIMIT 1
                ) exam_origin ON TRUE
                """
            )
        return " ".join(joins)

    def _problem_select_columns_sql(self, schema: Dict[str, object]) -> str:
        response_expr = self._response_select_expr(schema)
        course_expr = self._meta_select_expr(schema, "course_column", "COALESCE(t.area,'')", use_catalog=True)
        topic_expr = self._meta_select_expr(schema, "topic_column", "COALESCE(t.nombre,'')", use_catalog=True)
        subtopic_expr = self._meta_select_expr(schema, "subtopic_column", "COALESCE(s.nombre,'')", use_catalog=True)
        author_expr = self._meta_select_expr(schema, "author_column", "COALESCE(l.autor,'')", allow_fallback=True)
        editorial_expr = self._meta_select_expr(
            schema,
            "editorial_column",
            "COALESCE(l.editorial,'')",
            allow_fallback=True,
        )
        pdf_expr = "COALESCE(l.pdf_path,'')" if self._needs_book_join(schema) else "''"
        source_column = str(schema.get("source_file_column") or "")
        images_column = str(schema.get("images_column") or "")
        folder_column = str(schema.get("folder_column") or "")
        math_consistency_column = str(schema.get("math_consistency_column") or "")
        problem_type_column = str(schema.get("problem_type_column") or "")
        problem_cols = set(str(x) for x in schema.get("problem_columns", []))
        source_expr = f"COALESCE(p.{source_column},'')" if source_column else "''"
        exam_expr = "COALESCE(exam_origin.examen_nombre,'')" if schema.get("has_origins_table") else "''"
        return f"""
            p.id,
            p.numero_original,
            COALESCE(p.enunciado_latex,''),
            {response_expr},
            {course_expr},
            {topic_expr},
            {subtopic_expr},
            {author_expr},
            {editorial_expr},
            {pdf_expr},
            {source_expr},
            {f"COALESCE(p.{folder_column},'')" if folder_column else "''"},
            {f"COALESCE(p.{images_column}, ARRAY[]::text[])" if images_column else "ARRAY[]::text[]"},
            {f"COALESCE(p.{math_consistency_column},'')" if math_consistency_column else "'Sin revisar'"},
            {"COALESCE(p.codigo_instancia,'')" if "codigo_instancia" in problem_cols else "''"},
            {"COALESCE(p.libro_codigo,'')" if "libro_codigo" in problem_cols else "''"},
            {f"COALESCE(p.{problem_type_column},'')" if problem_type_column else "''"},
            {exam_expr}
        """

    def _problem_row_to_dict(self, row: tuple[object, ...]) -> Dict[str, object]:
        raw_images = row[12] if len(row) > 12 else []
        images: list[str] = []
        if isinstance(raw_images, (list, tuple)):
            for value in raw_images:
                clean = str(value or "").strip()
                if clean:
                    images.append(clean)
        respuesta = str(row[3] or "")
        item_text = str(row[2] or "")
        explicit_type = str(row[16] if len(row) > 16 else "")
        examen = str(row[17] if len(row) > 17 else "")
        return {
            "id": int(row[0]),
            "numero_original": int(row[1] or 0),
            "enunciado_latex": item_text,
            "respuesta_correcta": respuesta,
            "curso": self._canonical_course_value(row[4]),
            "tema": str(row[5] or ""),
            "subtema": str(row[6] or ""),
            "autor": str(row[7] or ""),
            "editorial": str(row[8] or ""),
            "pdf_path": str(row[9] or ""),
            "archivo_origen": str(row[10] or ""),
            "ruta_carpeta": str(row[11] or ""),
            "imagenes": images,
            "consistencia_matematica": str(row[13] or "") or "Sin revisar",
            "codigo_instancia": str(row[14] or ""),
            "libro_codigo": str(row[15] or ""),
            "tipo_problema": self._infer_problem_type(item_text, respuesta, explicit_type),
            "examen": examen,
        }

    def _natural_problem_order_sql(self, schema: Dict[str, object]) -> str:
        problem_cols = set(str(x) for x in schema.get("problem_columns", []))
        parts: list[str] = []
        if "libro_codigo" in problem_cols:
            parts.append("LOWER(COALESCE(p.libro_codigo, '')) ASC")
        if "codigo_instancia" in problem_cols:
            parts.append("LOWER(COALESCE(p.codigo_instancia, '')) ASC")
        parts.extend(["p.numero_original ASC", "p.id ASC"])
        return "ORDER BY " + ", ".join(parts)

    def _infer_problem_type(self, item_text: str, respuesta_correcta: str = "", explicit_type: str = "") -> str:
        raw_type = str(explicit_type or "").strip().lower().replace(" ", "_")
        if raw_type in {"abierto", "desarrollo", "respuesta_abierta"}:
            return "abierto"
        if raw_type in {"opcion_multiple", "opción_multiple", "multiple", "alternativas"}:
            return "opcion_multiple"
        if str(respuesta_correcta or "").strip():
            return "opcion_multiple"
        labels = {str(match.group(1) or "").upper() for match in OPTION_LABEL_RE.finditer(str(item_text or ""))}
        return "opcion_multiple" if {"A", "B"}.issubset(labels) else "abierto"

    def _build_filters(
        self,
        schema: Dict[str, object],
        *,
        curso: str,
        tema_id: Optional[object],
        subtema_id: Optional[object],
        autor: str,
        editorial: str,
        estado: str,
        clave: str,
    ) -> tuple[str, str, List[object]]:
        clauses: List[str] = []
        params: List[object] = []
        uses_catalog = bool(schema.get("uses_catalog"))
        join_sql = self._compose_join_sql(schema)
        topic_ref = self._parse_meta_ref(tema_id)
        subtopic_ref = self._parse_meta_ref(subtema_id)
        if uses_catalog:
            canonical_course = self._canonical_course_value(curso)
            if canonical_course:
                course_clauses: list[str] = []
                normalized_course = self._normalized_text_key(canonical_course)
                normalized_area_expr = self._sql_normalized_text("COALESCE(t.area,'')")
                course_clauses.append(f"{normalized_area_expr} = %s")
                params.append(normalized_course)
                if schema.get("course_column"):
                    normalized_course_expr = self._sql_normalized_text(f"p.{schema['course_column']}")
                    course_clauses.append(f"{normalized_course_expr} = %s")
                    params.append(normalized_course)
                clauses.append("(" + " OR ".join(course_clauses) + ")")

            topic_clauses: list[str] = []
            topic_text = str(topic_ref.get("text") or "").strip()
            if topic_ref.get("id") is not None and schema.get("tema_id_column"):
                topic_clauses.append("p.tema_id = %s")
                params.append(int(topic_ref["id"]))
            if topic_text:
                normalized_topic = self._normalized_text_key(topic_text)
                if schema.get("topic_column"):
                    normalized_topic_expr = self._sql_normalized_text(f"p.{schema['topic_column']}")
                    topic_clauses.append(f"{normalized_topic_expr} = %s")
                    params.append(normalized_topic)
                if schema.get("has_temas_table"):
                    normalized_catalog_topic_expr = self._sql_normalized_text("COALESCE(t.nombre,'')")
                    topic_clauses.append(f"{normalized_catalog_topic_expr} = %s")
                    params.append(normalized_topic)
            if topic_clauses:
                clauses.append("(" + " OR ".join(topic_clauses) + ")")

            subtopic_clauses: list[str] = []
            subtopic_text = str(subtopic_ref.get("text") or "").strip()
            if subtopic_ref.get("id") is not None and schema.get("subtema_id_column"):
                subtopic_clauses.append("p.subtema_id = %s")
                params.append(int(subtopic_ref["id"]))
            if subtopic_text:
                normalized_subtopic = self._normalized_text_key(subtopic_text)
                if schema.get("subtopic_column"):
                    normalized_subtopic_expr = self._sql_normalized_text(f"p.{schema['subtopic_column']}")
                    subtopic_clauses.append(f"{normalized_subtopic_expr} = %s")
                    params.append(normalized_subtopic)
                if schema.get("has_subtemas_table"):
                    normalized_catalog_subtopic_expr = self._sql_normalized_text("COALESCE(s.nombre,'')")
                    subtopic_clauses.append(f"{normalized_catalog_subtopic_expr} = %s")
                    params.append(normalized_subtopic)
            if subtopic_clauses:
                clauses.append("(" + " OR ".join(subtopic_clauses) + ")")
        else:
            canonical_course = self._canonical_course_value(curso)
            if canonical_course and schema.get("course_column"):
                normalized_course_expr = self._sql_normalized_text(f"p.{schema['course_column']}")
                clauses.append(f"{normalized_course_expr} = %s")
                params.append(self._normalized_text_key(canonical_course))
            topic_text = str(topic_ref.get("text") or tema_id or "").strip()
            if topic_text and schema.get("topic_column"):
                normalized_topic_expr = self._sql_normalized_text(f"p.{schema['topic_column']}")
                clauses.append(f"{normalized_topic_expr} = %s")
                params.append(self._normalized_text_key(topic_text))
            subtopic_text = str(subtopic_ref.get("text") or subtema_id or "").strip()
            if subtopic_text and schema.get("subtopic_column"):
                normalized_subtopic_expr = self._sql_normalized_text(f"p.{schema['subtopic_column']}")
                clauses.append(f"{normalized_subtopic_expr} = %s")
                params.append(self._normalized_text_key(subtopic_text))

        author_expr = self._filter_field_expr(schema, "author_column")
        editorial_expr = self._filter_field_expr(schema, "editorial_column")
        if (autor or "").strip() and author_expr:
            clauses.append(f"{self._sql_normalized_text(author_expr)} = %s")
            params.append(self._normalized_text_key(autor))
        if (editorial or "").strip() and editorial_expr:
            clauses.append(f"{self._sql_normalized_text(editorial_expr)} = %s")
            params.append(self._normalized_text_key(editorial))

        math_consistency_column = str(schema.get("math_consistency_column") or "")
        estado = (estado or "").strip()
        if math_consistency_column:
            if estado in {"Pendiente Revision", "Sin revisar"}:
                clauses.append(
                    f"(p.{math_consistency_column} IS NULL OR p.{math_consistency_column} = 'Sin revisar')"
                )
            elif estado in {"Bien Planteado", "Consistente"}:
                clauses.append(f"p.{math_consistency_column} = %s")
                params.append("Consistente")
            elif estado in {"Mal Planteado", "Inconsistente"}:
                clauses.append(f"p.{math_consistency_column} = %s")
                params.append("Inconsistente")

        clave = (clave or "").strip()
        response_column = str(schema.get("response_column") or "")
        key_flag_column = str(schema.get("key_flag_column") or "")
        problem_type_column = str(schema.get("problem_type_column") or "")
        has_response = bool(response_column)
        has_key_flag = bool(key_flag_column)
        response_has_key_sql = (
            f"COALESCE(TRIM(CAST(p.{response_column} AS text)), '') <> ''"
            if has_response
            else "FALSE"
        )
        key_flag_sql = f"COALESCE(p.{key_flag_column}, FALSE)" if has_key_flag else "FALSE"
        has_key_sql = f"({response_has_key_sql} OR {key_flag_sql})"
        explicit_type_sql = f"LOWER(COALESCE(TRIM(CAST(p.{problem_type_column} AS text)), ''))" if problem_type_column else "''"
        looks_multiple_sql = (
            f"({explicit_type_sql} IN ('opcion_multiple', 'opción_multiple', 'multiple', 'alternativas') "
            f"OR ({explicit_type_sql} = '' AND ({has_key_sql} OR "
            "((COALESCE(p.enunciado_latex,'') ~* '(^|[^A-Za-z0-9])A\\)') "
            "AND (COALESCE(p.enunciado_latex,'') ~* '(^|[^A-Za-z0-9])B\\)')))))"
        )
        looks_open_sql = (
            f"({explicit_type_sql} IN ('abierto', 'desarrollo', 'respuesta_abierta') "
            f"OR ({explicit_type_sql} = '' AND NOT {looks_multiple_sql}))"
        )
        if clave == "Con clave":
            clauses.append(has_key_sql)
        elif clave == "Sin clave":
            clauses.append(f"({looks_multiple_sql} AND NOT {has_key_sql})")
        elif clave in {"Abiertos", "Abierto"}:
            clauses.append(looks_open_sql)

        where_sql = ""
        if clauses:
            where_sql = "WHERE " + " AND ".join(f"({c})" for c in clauses)
        return join_sql, where_sql, params

    def _limpiar_enunciado_para_word(self, enunciado: str) -> str:
        text = TAG_META_RE.sub(" ", enunciado or "")
        text = text.replace("\u00a3", "\n")
        text = text.replace("\u00e6", "\n")
        text = text.replace("\u00c2\u00a3", "\n")
        text = text.replace("\u00c3\u00a6", "\n")
        text = re.sub(r"[ \t]{2,}", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _normalizar_item_para_latex(self, raw_item: str, fallback_idx: int) -> str:
        src = str(raw_item or "").replace("\r\n", "\n").replace("\r", "\n").strip()
        if not src:
            return ""

        src = TAG_META_RE.sub(" ", src)
        src = IMG_MARKER_RE.sub(" ", src)
        src = src.replace("\u00a3", "\n").replace("\u00e6", "\n")
        src = src.replace("\u00c2\u00a3", "\n").replace("\u00c3\u00a6", "\n")
        src = re.sub(r"[ \t]{2,}", " ", src)
        src = re.sub(r"\n{3,}", "\n\n", src).strip()

        m = ITEM_HEADER_RE.match(src)
        if m:
            head = m.group(1).strip()
            body = m.group(2).strip()
            if body:
                body = body.replace("\n", "\n\\\\ ")
                return f"{head} {body}".strip()
            return head

        plain = src.replace("\n", "\n\\\\ ").strip()
        return f"\\item[\\textbf{{{max(1, int(fallback_idx))}.}}] {plain}"

    def _escape_latex_text(self, text: str) -> str:
        value = str(text or "")
        value = value.replace("\\", "\\textbackslash{}")
        value = value.replace("&", "\\&")
        value = value.replace("%", "\\%")
        value = value.replace("#", "\\#")
        value = value.replace("_", "\\_")
        value = value.replace("{", "\\{")
        value = value.replace("}", "\\}")
        value = value.replace("$", "\\$")
        value = value.replace("^", "\\textasciicircum{}")
        value = value.replace("~", "\\textasciitilde{}")
        return value

    def _ensure_enumerate_wrapper(self, source_text: str) -> str:
        text = str(source_text or "").strip()
        if not text:
            return "\\begin{enumerate}\n\\end{enumerate}"
        low = text.lower()
        if "\\begin{enumerate}" in low and "\\end{enumerate}" in low:
            return text
        return "\\begin{enumerate}\n" + text + "\n\\end{enumerate}"
